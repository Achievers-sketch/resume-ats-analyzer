import io
import json
import os
import re
from typing import Any, Dict, List

import streamlit as st
from docx import Document
from google import genai
from google.genai import types
from pypdf import PdfReader


MODEL_NAME = "gemini-2.5-flash"
MAX_RESUME_CHARS = 30000


st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)


def get_api_key() -> str:
    """Read Gemini API key from Streamlit secrets first, then environment variables."""
    try:
        key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        key = ""

    return key or os.getenv("GEMINI_API_KEY", "")


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))

    pages = []

    for page in reader.pages:
        pages.append(page.extract_text() or "")

    return "\n".join(pages)


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))

    parts = [
        p.text
        for p in document.paragraphs
        if p.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            parts.append(
                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )
            )

    return "\n".join(parts)


def extract_resume_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()

    extension = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if extension == "pdf":
        text = extract_pdf_text(data)

    elif extension == "docx":
        text = extract_docx_text(data)

    else:
        raise ValueError(
            "Unsupported file type. Please upload a PDF or DOCX resume."
        )

    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return text[:MAX_RESUME_CHARS]


def local_resume_checks(resume: str) -> Dict[str, Any]:
    """
    Deterministic baseline checks.

    This gives the application some objective resume checks
    instead of relying entirely on the LLM.
    """

    lower = resume.lower()

    word_count = len(
        re.findall(
            r"\b\w+[\w+.#/-]*\b",
            resume
        )
    )

    sections = {
        "Contact information": bool(
            re.search(
                r"(?:email|@|phone|tel|linkedin|github)",
                lower
            )
        ),

        "Summary / Profile": bool(
            re.search(
                r"\b(summary|profile|objective|professional summary)\b",
                lower
            )
        ),

        "Experience": bool(
            re.search(
                r"\b(experience|employment|work history|professional experience)\b",
                lower
            )
        ),

        "Education": bool(
            re.search(
                r"\b(education|academic|degree|university|college)\b",
                lower
            )
        ),

        "Skills": bool(
            re.search(
                r"\b(skills|technical skills|core competencies)\b",
                lower
            )
        ),
    }

    action_verbs = len(
        re.findall(
            r"\b("
            r"achieved|built|created|developed|designed|implemented|"
            r"improved|increased|led|managed|optimized|reduced|"
            r"launched|automated|delivered|analyzed|engineered"
            r")\b",
            lower,
        )
    )

    quantified = len(
        re.findall(
            r"(?:"
            r"\b\d+(?:\.\d+)?%"
            r"|\$\s?\d+[\d,]*"
            r"|\b\d+[\d,]*\+?\b"
            r")",
            resume,
        )
    )

    bullets = len(
        re.findall(
            r"(?:^|\n)\s*[•●▪◦\-*]\s+",
            resume
        )
    )

    score = 0

    score += 30 if sections["Contact information"] else 0
    score += 15 if sections["Experience"] else 0
    score += 15 if sections["Education"] else 0
    score += 15 if sections["Skills"] else 0
    score += 10 if sections["Summary / Profile"] else 0
    score += 5 if word_count >= 250 else 0
    score += 5 if action_verbs >= 5 else 0
    score += 5 if quantified >= 3 else 0

    return {
        "baseline_score": min(score, 100),
        "word_count": word_count,
        "action_verbs": action_verbs,
        "quantified_items": quantified,
        "bullet_count": bullets,
        "sections": sections,
    }


RESPONSE_SCHEMA = {
    "type": "object",

    "properties": {
        "ats_score": {
            "type": "integer",
            "minimum": 0,
            "maximum": 100,
        },

        "score_label": {
            "type": "string",
        },

        "summary": {
            "type": "string",
        },

        "strengths": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "improvements": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "missing_keywords": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "formatting_issues": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "section_feedback": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "rewritten_bullets": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },

        "action_plan": {
            "type": "array",
            "items": {
                "type": "string"
            },
        },
    },

    "required": [
        "ats_score",
        "score_label",
        "summary",
        "strengths",
        "improvements",
        "missing_keywords",
        "formatting_issues",
        "section_feedback",
        "rewritten_bullets",
        "action_plan",
    ],
}


def analyze_with_gemini(
    resume: str,
    job_description: str,
    baseline: Dict[str, Any],
) -> Dict[str, Any]:

    api_key = get_api_key()

    if not api_key:
        raise RuntimeError(
            "Gemini API key not found. "
            "Add GEMINI_API_KEY to Streamlit secrets "
            "or your environment."
        )

    client = genai.Client(api_key=api_key)

    jd = (
        job_description.strip()
        or
        "No job description supplied. "
        "Evaluate general ATS readiness rather than "
        "job-specific matching."
    )

    prompt = f"""
You are an expert ATS resume evaluator and career coach.

Evaluate the resume below.

Give an ATS-readiness score from 0 to 100.

Use these scoring criteria:

- Keyword/job-description alignment:
  35 points when a job description is provided.
  Otherwise evaluate general relevance.

- ATS-friendly structure and section headings:
  20 points.

- Measurable achievements and strong action verbs:
  15 points.

- Clarity and concise professional language:
  15 points.

- Formatting/parser safety based only on the extracted text:
  15 points.

Important rules:

1. Do not claim that you can see visual formatting that
   was not present in the extracted text.

2. Flag likely formatting issues only when supported
   by the extracted text.

3. Do not invent experience, qualifications,
   employers, dates, or skills.

4. Suggested rewrites must preserve the user's facts.

5. If a metric is missing, use a placeholder such as
   [X%] rather than inventing a number.

6. Missing keywords should be relevant to the supplied
   job description.

Optional target job description:

---
{jd[:15000]}
---

Deterministic text checks:

{json.dumps(baseline, indent=2)}

Resume text:

---
{resume}
---

Return only JSON matching the requested schema.
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
            response_schema=RESPONSE_SCHEMA,
        ),
    )

    if not response.text:
        raise RuntimeError(
            "Gemini returned an empty response."
        )

    return json.loads(response.text)


def score_color(score: int) -> str:

    if score >= 80:
        return "Excellent"

    if score >= 65:
        return "Good"

    if score >= 50:
        return "Needs improvement"

    return "Needs major improvement"


def show_list(
    title: str,
    items: List[str],
    icon: str = "•",
) -> None:

    st.subheader(title)

    if items:

        for item in items:
            st.markdown(
                f"{icon} {item}"
            )

    else:
        st.caption(
            "No items identified."
        )


st.title("📄 Resume ATS Analyzer")

st.caption(
    "Upload your resume, optionally add a target job description, "
    "and get an ATS score plus actionable improvements."
)


with st.sidebar:

    st.header("How it works")

    st.write(
        "1. Upload a PDF or DOCX resume."
    )

    st.write(
        "2. Optionally paste the job description "
        "for targeted keyword matching."
    )

    st.write(
        "3. Gemini evaluates ATS readiness "
        "and gives improvement suggestions."
    )

    st.divider()

    st.caption(
        "Files are processed for the current analysis "
        "and are not saved by this app."
    )


uploaded = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx"],
    help="PDF or DOCX only. Text-based PDFs work best.",
)


job_description = st.text_area(
    "Target job description (optional)",
    height=220,
    placeholder=(
        "Paste the job description here for a more useful "
        "ATS keyword match..."
    ),
)


analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)


if analyze_button:

    if not uploaded:

        st.warning(
            "Please upload a PDF or DOCX resume first."
        )

        st.stop()

    try:

        with st.spinner(
            "Extracting resume text..."
        ):

            resume_text = extract_resume_text(
                uploaded
            )

            if len(resume_text.strip()) < 80:

                st.error(
                    "Very little text could be extracted. "
                    "If this is a scanned/image-only PDF, "
                    "use a text-based PDF or DOCX."
                )

                st.stop()

            baseline = local_resume_checks(
                resume_text
            )

        with st.spinner(
            "Gemini is analyzing your resume..."
        ):

            result = analyze_with_gemini(
                resume_text,
                job_description,
                baseline,
            )

        score = int(
            max(
                0,
                min(
                    100,
                    result.get(
                        "ats_score",
                        baseline["baseline_score"]
                    ),
                ),
            )
        )

        st.session_state["analysis"] = result
        st.session_state["score"] = score
        st.session_state["baseline"] = baseline
        st.session_state["resume_text"] = resume_text

    except Exception as exc:

        st.error(
            f"Analysis failed: {exc}"
        )

        st.info(
            "Check your Gemini API key, internet connection, "
            "and that the uploaded file contains selectable text."
        )


if "analysis" in st.session_state:

    result = st.session_state["analysis"]

    score = st.session_state["score"]

    baseline = st.session_state["baseline"]

    st.divider()

    st.subheader("ATS Score")

    c1, c2, c3, c4 = st.columns(4)

    c1.metric(
        "ATS score",
        f"{score}/100"
    )

    c2.metric(
        "Score level",
        result.get(
            "score_label",
            score_color(score)
        )
    )

    c3.metric(
        "Words",
        baseline["word_count"]
    )

    c4.metric(
        "Quantified items",
        baseline["quantified_items"]
    )

    st.progress(
        score / 100
    )

    st.write(
        result.get(
            "summary",
            ""
        )
    )

    left, right = st.columns(2)

    with left:

        show_list(
            "Strengths",
            result.get("strengths", []),
            "✅",
        )

        show_list(
            "Improvements",
            result.get("improvements", []),
            "🔧",
        )

        show_list(
            "Missing keywords",
            result.get("missing_keywords", []),
            "🔑",
        )

    with right:

        show_list(
            "Formatting / ATS issues",
            result.get("formatting_issues", []),
            "⚠️",
        )

        show_list(
            "Section feedback",
            result.get("section_feedback", []),
            "📌",
        )

        show_list(
            "Action plan",
            result.get("action_plan", []),
            "➡️",
        )

    st.subheader(
        "Suggested bullet improvements"
    )

    for bullet in result.get(
        "rewritten_bullets",
        []
    ):

        st.markdown(
            f"- {bullet}"
        )

    with st.expander(
        "View extracted resume text"
    ):

        st.text(
            st.session_state["resume_text"]
        )
